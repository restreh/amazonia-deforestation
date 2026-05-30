"""Construye el documento final del proyecto integrador en formato .docx.

Uso:
    python scripts/build_final_delivery.py

Estructura siguiendo el template proporcionado por la maestria:
portada, introduccion, marco teorico, desarrollo metodologico de ML,
tecnologia (Big Data), visualizacion y comunicacion, conclusiones y
referencias. Las cifras y resultados se toman directamente de los
artefactos del proyecto (eval_*.json, mcnemar.json, bootstrap_spatial.json,
concordance_lin.json, compare_cv.json, benchmarks/).

Dependencias: python-docx.
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Cargadores de artefactos del proyecto
# ---------------------------------------------------------------------------
INTERIM = ROOT / "data" / "interim"


def _load(name: str) -> dict | None:
    p = INTERIM / name
    if not p.exists():
        return None
    return json.loads(p.read_text(encoding="utf-8"))


def _safe_get(d, *keys, default=None):
    cur = d
    for k in keys:
        if cur is None or k not in cur:
            return default
        cur = cur[k]
    return cur


EVAL_XGB = _load("eval_xgboost.json") or {}
EVAL_RF = _load("eval_random_forest.json") or {}
EVAL_UNET = _load("eval_unet.json") or {}
EVAL_ENS = _load("eval_ensemble.json") or {}
EVAL_IMG = _load("eval_unet_imagenet.json") or {}
MCNEMAR = _load("mcnemar.json") or {}
BOOT = _load("bootstrap_spatial.json") or {}
CCC = _load("concordance_lin.json") or {}
COMP_CV = _load("compare_cv.json") or {}
CV_SUM = _load("baseline_cv_summary.json") or {}


# ---------------------------------------------------------------------------
# Helpers de estilo
# ---------------------------------------------------------------------------
def set_body_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Calibri")


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_par(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            bold=False, italic=False, size: int | None = None,
            space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_centered(doc: Document, text: str, bold: bool = False,
                 italic: bool = False, size: int | None = None,
                 space_after: int = 6):
    return add_par(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold,
                   italic=italic, size=size, space_after=space_after)


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table_simple(doc: Document, headers: list[str], rows: list[list[str]],
                     header_bold: bool = True, col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = header_bold
        run.font.size = Pt(10)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    return table


def add_reference(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text).font.size = Pt(10)


def fmt_decimal(x, n=3):
    if x is None:
        return "—"
    return f"{x:.{n}f}"


def fmt_pct(x, n=2):
    if x is None:
        return "—"
    return f"{x*100:.{n}f} %"


# ---------------------------------------------------------------------------
# Secciones del documento
# ---------------------------------------------------------------------------
def section_portada(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "MAESTRÍA EN CIENCIA DE DATOS Y ANALÍTICA", bold=True, size=12)
    add_centered(doc, "Proyecto Integrado 1 — 2026-1", size=11, space_after=24)
    add_centered(doc,
                 "Detección temprana de deforestación en la Amazonía colombiana "
                 "mediante Sentinel-2 y aprendizaje automático",
                 bold=True, size=18, space_after=24)
    add_centered(doc, "Documento final de entrega", italic=True, size=12,
                 space_after=18)

    add_centered(doc, "Equipo CLRCV", bold=True, size=12, space_after=4)
    integrantes = [
        "Gia Mariana Calle Higuita",
        "Juan Diego Llorente Ortega",
        "Juan José Restrepo Higuita",
        "Manuela Caro Villada",
        "Jerónimo Velásquez Escobar",
    ]
    for nombre in integrantes:
        add_centered(doc, nombre, size=11, space_after=2)
    doc.add_paragraph()

    add_centered(doc, "Universidad EAFIT", size=11, space_after=2)
    add_centered(doc, f"{date.today().strftime('%d de mayo de 2026')}",
                 size=11, space_after=2)
    add_centered(doc,
                 "Repositorio: https://github.com/restreh/amazonia-deforestation",
                 size=10)
    add_centered(doc,
                 "Tablero público: https://amazonia-deforestation.streamlit.app/",
                 size=10)
    doc.add_page_break()


def section_introduccion(doc: Document):
    add_heading(doc, "1. Introducción", level=1)
    add_par(doc,
        "Según el Boletín 42 de Detección Temprana de Deforestación del Sistema "
        "de Monitoreo de Bosques y Carbono (SMByC) del Instituto de Hidrología, "
        "Meteorología y Estudios Ambientales (IDEAM), durante 2024 se "
        "deforestaron aproximadamente 107.000 hectáreas en Colombia, un 35 % "
        "más que las 79.000 hectáreas registradas en 2023 (IDEAM, 2025). La "
        "región amazónica concentra históricamente la mayor proporción de "
        "pérdida, con los departamentos de Caquetá, Meta, Guaviare y Putumayo "
        "en las primeras posiciones nacionales. El Boletín 45, correspondiente "
        "al cuarto trimestre de 2025, reporta 36.129 hectáreas con alertas "
        "tempranas en la Amazonía colombiana entre octubre y diciembre, un "
        "incremento de 16.330 hectáreas respecto al mismo trimestre de 2024 "
        "(IDEAM, 2026). Caquetá lidera el incremento con +9.078 hectáreas y "
        "concentra cuatro de los 21 núcleos activos identificados a nivel "
        "amazónico, entre los que destaca Cuemaní, ubicado entre los "
        "municipios de San Vicente del Caguán y Cartagena del Chairá.")
    add_par(doc,
        "El SMByC publica boletines trimestrales basados en Landsat, Sentinel-1 "
        "y Sentinel-2. La latencia trimestral limita la capacidad de respuesta "
        "operativa de las autoridades ambientales, Parques Nacionales Naturales "
        "y organizaciones de la sociedad civil frente a núcleos activos en "
        "evolución. El proyecto aborda este problema desarrollando un sistema "
        "académico complementario, con código de acceso abierto y un ciclo de "
        "actualización potencialmente más corto, que detecta polígonos de "
        "deforestación en un área de interés del arco amazónico colombiano a "
        "partir de composiciones trimestrales de Sentinel-2 y aprendizaje "
        "automático supervisado, utilizando Hansen Global Forest Change (Hansen "
        "et al., 2013) como referencia de entrenamiento y evaluación.")
    add_par(doc,
        "La pregunta de investigación que orienta el trabajo es la siguiente: "
        "¿con qué desempeño puede un modelo supervisado, entrenado sobre "
        "composiciones trimestrales de Sentinel-2, detectar polígonos de "
        "deforestación en un núcleo activo del Caquetá usando Hansen Global "
        "Forest Change como referencia, bajo un protocolo de validación que "
        "controla la dependencia espacial? Pacheco-Pascagaza et al. (2022) "
        "habían reportado para México y Colombia que modelos de aprendizaje "
        "automático aplicados a Sentinel-2 alcanzan niveles de concordancia "
        "útiles frente a las bases oficiales; el proyecto valida esa "
        "afirmación en una ventana específica del Caquetá, agrega un protocolo "
        "estadístico de comparación más estricto que el habitual en la "
        "literatura y publica un tablero interactivo que comunica los "
        "resultados a usuarios técnicos y operativos.")
    add_par(doc,
        "El proyecto articula tres materias de la maestría. Aprendizaje "
        "Automático aporta los modelos supervisados y el protocolo estadístico "
        "de evaluación. Almacenamiento y Procesamiento de Grandes Datos aporta "
        "la arquitectura en AWS: ingestión vía STAC y Cloud-Optimized GeoTIFF, "
        "almacenamiento en Amazon S3 con particionamiento Hive, consultas en "
        "Amazon Athena y servicio de inferencia en AWS Lambda. Visualización "
        "de Datos aporta el tablero interactivo en Streamlit, desplegado "
        "públicamente, que presenta el contexto del problema, el diagnóstico "
        "espacial, la comparación entre modelos, el mapa de predicciones, el "
        "análisis municipal y el detalle del despliegue Big Data.")
    add_par(doc,
        "El alcance espacial se limita a una ventana de 5.023 km² sobre el "
        "núcleo activo de Cuemaní, dentro de los municipios de Cartagena del "
        "Chairá y San Vicente del Caguán. El alcance temporal corresponde al "
        "año calendario 2024, agregado en cuatro composiciones trimestrales. "
        "El alcance metodológico cubre cuatro modelos candidatos (Random "
        "Forest, XGBoost, U-Net entrenada desde cero y ensamble por promedio "
        "ponderado XGBoost + U-Net) y un experimento de control con encoder "
        "preentrenado en ImageNet, evaluados con métricas píxel y polígono, "
        "prueba de McNemar pareada, bootstrap espacial por bloques con B = "
        "1.000 y coeficiente de concordancia de Lin.")
    doc.add_page_break()


def section_marco_teorico(doc: Document):
    add_heading(doc, "2. Marco teórico y referencias", level=1)
    add_par(doc,
        "El monitoreo de cobertura boscosa con sensores ópticos de media "
        "resolución se apoya en la diferenciación espectral entre superficies "
        "vegetales activas y superficies alteradas. Sentinel-2 opera con 13 "
        "bandas espectrales, resolución espacial de 10 a 60 metros y revisita "
        "combinada de 5 días para la constelación Sentinel-2A y Sentinel-2B "
        "(Drusch et al., 2012), lo que permite construir composiciones "
        "temporales que mitigan la presencia de nubes en zonas tropicales. "
        "Los índices derivados a partir de bandas del visible, infrarrojo "
        "cercano e infrarrojo de onda corta, como el Normalized Difference "
        "Vegetation Index (NDVI), el Normalized Burn Ratio (NBR) y el "
        "Normalized Difference Water Index (NDWI), funcionan como proxies de "
        "la pérdida de cobertura y son entradas estándar en sistemas de "
        "detección de deforestación con teledetección óptica.")

    add_heading(doc, "2.1 Enfoques clásicos y profundos", level=2)
    add_par(doc,
        "Los enfoques clásicos de clasificación supervisada se basan en "
        "Random Forest (Breiman, 2001) y Gradient Boosting (Chen & Guestrin, "
        "2016) sobre vectores de atributos por píxel derivados de bandas "
        "espectrales e índices, con la opción de incorporar contexto espacial "
        "mediante atributos calculados sobre ventanas locales. Estos modelos "
        "ofrecen tiempos de ajuste compatibles con infraestructura de nube con "
        "restricciones de cómputo y soportan análisis de importancia de "
        "atributos directamente interpretable.")
    add_par(doc,
        "Los enfoques basados en redes convolucionales aplican arquitecturas "
        "tipo encoder-decoder sobre ventanas espaciales. La U-Net "
        "(Ronneberger et al., 2015), originalmente desarrollada para "
        "segmentación biomédica, se ha adaptado a teledetección por su "
        "capacidad de preservar resolución espacial mediante skip connections "
        "y su buen comportamiento con conjuntos de entrenamiento moderados. "
        "Para deforestación amazónica, Adarme et al. (2022) evaluaron variantes "
        "de redes totalmente convolucionales sobre Landsat con etiquetas "
        "PRODES y reportaron F1 píxel entre 0,45 y 0,60. Maretto et al. (2021) "
        "combinaron U-Net con módulos recurrentes para capturar dinámica "
        "temporal y reportaron F1 píxel cercano a 0,62 sobre el mismo "
        "dominio. Trabajos más recientes han explorado arquitecturas basadas "
        "en transformers y modelos fundacionales preentrenados sobre imágenes "
        "satelitales, como Prithvi-EO-2.0 (Szwarcman et al., 2024), que "
        "reportan mejoras en regímenes con pocas etiquetas pero que exceden "
        "el alcance computacional del presente proyecto.")
    add_par(doc,
        "Para el contexto colombiano, Pacheco-Pascagaza et al. (2022) "
        "propusieron un sistema de detección de cambios casi en tiempo real "
        "con Sentinel-2 y aprendizaje automático sobre bosques mexicanos y "
        "colombianos, reportando concordancias útiles frente a productos "
        "oficiales. Hansen et al. (2013) publicaron el producto global de "
        "cambio en cobertura forestal a 30 metros, actualizado anualmente, "
        "que se ha consolidado como referencia de evaluación y fuente de "
        "etiquetas débiles. Armenteras et al. (2021) analizaron la pérdida "
        "de biodiversidad asociada a incendios y deforestación en América "
        "Latina con datos de teledetección, contextualizando la magnitud "
        "ambiental del problema.")

    add_heading(doc, "2.2 Dependencia espacial y su tratamiento estadístico", level=2)
    add_par(doc,
        "Las observaciones derivadas de teledetección presentan dependencia "
        "espacial. Píxeles próximos tienden a tener valores más similares "
        "que píxeles distantes, lo que constituye una manifestación de la "
        "primera ley de Tobler. Esta propiedad se cuantifica mediante el "
        "índice I de Moran (Moran, 1950) y el semivariograma empírico, y "
        "tiene dos consecuencias para el modelado.")
    add_par(doc,
        "Primero, la dependencia entre vecinos es información útil. Puede "
        "aprovecharse en los baselines mediante atributos contextuales "
        "calculados sobre ventanas locales (medias, desviaciones, métricas "
        "GLCM) y en arquitecturas convolucionales mediante el campo receptivo "
        "creciente del encoder. Segundo, si los conjuntos de entrenamiento y "
        "prueba se construyen por muestreo aleatorio simple sobre la misma "
        "escena, la cercanía espacial entre ambos genera fuga de información "
        "y produce estimaciones de desempeño optimistas que no se sostienen "
        "al aplicar el modelo en zonas no observadas durante el entrenamiento "
        "(Roberts et al., 2017; Ploton et al., 2020; Karasiak et al., 2022). "
        "El procedimiento honesto, adoptado en este proyecto, consiste en "
        "particionar el área de interés en bloques espaciales de tamaño "
        "mayor o igual al rango del semivariograma y reportar las métricas "
        "bajo dos esquemas en paralelo: validación cruzada aleatoria como "
        "referencia comparable con la literatura previa que no controla por "
        "dependencia, y validación cruzada espacial por bloques como "
        "estimación honesta del desempeño esperado en zonas nuevas. La "
        "diferencia entre ambas estimaciones cuantifica el sesgo optimista "
        "introducido por la dependencia espacial y, por sí misma, constituye "
        "un resultado del proyecto.")

    add_heading(doc, "2.3 Métricas y comparación estadística", level=2)
    add_par(doc,
        "Bajo desbalance severo de clases (prevalencia positiva inferior al "
        "3 % en el área de interés), la exactitud global es engañosa: un "
        "clasificador trivial que prediga la clase mayoritaria alcanza una "
        "exactitud superior al 98 % sin detectar deforestación. Por esa "
        "razón se reportan precisión, recall, F1 e Intersección sobre Unión "
        "(IoU) sobre la clase positiva, junto con el área bajo la curva "
        "precisión-recall (AUC-PR), preferible al AUC-ROC bajo "
        "desbalance porque concentra la información en la región de "
        "interés. La métrica primaria es F1 sobre la clase positiva.")
    add_par(doc,
        "La comparación pareada entre modelos se realiza mediante la prueba "
        "de McNemar, que opera sobre los aciertos discordantes de cada par "
        "en el mismo conjunto de prueba y reporta un p-valor exacto vía "
        "distribución binomial. Para acotar el sesgo estadístico inducido por "
        "la autocorrelación espacial se complementa con un bootstrap por "
        "bloques (B = 1.000 iteraciones, remuestreo con reposición de los "
        "bloques de prueba), que produce intervalos de confianza al 95 % de "
        "F1, IoU y AUC-PR coherentes con la estructura espacial de los datos. "
        "Adicionalmente, el coeficiente de concordancia de Lin (Lin, 1989) "
        "evalúa la concordancia entre las hectáreas detectadas por bloque "
        "(modelo) y por bloque (Hansen), capturando simultáneamente exactitud "
        "y precisión en una sola métrica entre 0 y 1.")

    add_heading(doc, "2.4 Marco metodológico", level=2)
    add_par(doc,
        "El proyecto sigue la metodología CRISP-DM (Chapman et al., 2000), "
        "organizada en seis fases: entendimiento del negocio, entendimiento "
        "de los datos, preparación de los datos, modelado, evaluación y "
        "despliegue. La arquitectura técnica se apoya en especificaciones "
        "abiertas: SpatioTemporal Asset Catalog (STAC) para indexación de "
        "imágenes, Cloud-Optimized GeoTIFF (COG) para lectura por ventanas "
        "sin descarga completa y Apache Parquet con compresión Snappy para "
        "almacenamiento columnar de datos derivados. Los principios de "
        "diseño visual del tablero siguen las recomendaciones de Tufte (2001) "
        "sobre ratio dato-tinta y minimización de elementos no informativos.")
    doc.add_page_break()


def section_metodologia_ml(doc: Document):
    add_heading(doc, "3. Desarrollo metodológico de modelos de ML", level=1)

    add_heading(doc, "3.1 Entendimiento del problema", level=2)
    add_par(doc,
        "El problema se formaliza como una tarea de clasificación binaria "
        "a nivel de píxel. Sea Ω ⊂ ℝ² el área de interés discretizada en una "
        "grilla regular indexada por s. Para cada píxel y cada año t se "
        "observa un vector de atributos x(s, t) ∈ ℝᵈ derivado de "
        "composiciones trimestrales de Sentinel-2 (bandas espectrales e "
        "índices, opcionalmente atributos contextuales), y una etiqueta "
        "y(s, t) ∈ {0, 1}, donde 1 indica pérdida de cobertura arbórea según "
        "Hansen Global Forest Change.")
    add_par(doc,
        "El objetivo es aprender una función de decisión f: ℝᵈ → [0, 1] tal "
        "que ŷ(s, t) = 1 si f(x(s, t)) ≥ τ y 0 en caso contrario, donde τ es "
        "un umbral calibrado en el conjunto de validación a la prevalencia "
        "real, maximizando F1. La estimación de parámetros se realiza por "
        "minimización del riesgo empírico regularizado sobre el conjunto de "
        "entrenamiento. Para los baselines tabulares (Random Forest y "
        "XGBoost), x(s, t) agrega información espectral y contextual por "
        "píxel; para la U-Net, la entrada es un tensor de 56 canales sobre "
        "una ventana 256 × 256 y la salida es un mapa de probabilidades de "
        "la misma resolución.")

    add_heading(doc, "3.2 Análisis Exploratorio de Datos", level=2)
    add_heading(doc, "3.2.1 Entendimiento de los datos", level=3)
    add_par(doc,
        "Las fuentes principales de datos son Sentinel-2 Level-2A (Drusch "
        "et al., 2012) y Hansen Global Forest Change v1.12 (Hansen et al., "
        "2013). Las imágenes Sentinel-2 se acceden vía AWS Open Data Registry "
        "sobre el bucket público s3://sentinel-cogs en la región us-west-2 "
        "del proveedor, indexadas mediante la API STAC Earth-Search. El "
        "producto Hansen GFC v1.12 cubre el periodo 2000-2024 con resolución "
        "30 metros y se obtuvo desde el portal del Global Land Analysis and "
        "Discovery Lab. Como referencia adicional se incorporan los "
        "boletines trimestrales del SMByC-IDEAM (2025, 2026) y los límites "
        "municipales del Instituto Geográfico Agustín Codazzi.")
    add_par(doc,
        "El área de interés se delimitó a una ventana de 5.023 km² dentro "
        "de los municipios de Cartagena del Chairá y San Vicente del Caguán, "
        "Caquetá, sobre el núcleo activo Cuemaní. La selección utilizó la "
        "capa lossyear de Hansen para identificar las celdas con mayor "
        "concentración de pérdida en 2024 y maximizar la cobertura del "
        "fenómeno: el AOI captura aproximadamente el 54 % de la pérdida "
        "municipal 2024 reportada por Hansen para ese par de municipios.")

    add_heading(doc, "3.2.2 Preparación de los datos", level=3)
    add_par(doc,
        "Las imágenes Sentinel-2 se filtran por nubosidad inferior al 80 % "
        "y se mascarillan a nivel de píxel con la capa SCL (clase de nube, "
        "sombra de nube, nieve). Las composiciones trimestrales se generan "
        "por agregación pixel-wise sobre cada trimestre del año objetivo, "
        "usando dos estadísticos por píxel: la mediana (resistente a "
        "outliers atmosféricos residuales) y el percentil 25 (sensible a "
        "superficies oscuras o quemadas tras la pérdida de cobertura). Las "
        "composiciones se almacenan como Cloud-Optimized GeoTIFF a "
        "resolución de trabajo de 20 metros, reproyectando las bandas de 10 "
        "metros por promedio.")
    add_par(doc,
        "Los índices NDVI, NBR y NDWI se calculan por trimestre sobre la "
        "composición mediana. La etiqueta de Hansen para el año objetivo "
        "(lossyear == 24, correspondiente a 2024) se realinea a 20 metros "
        "por vecino más cercano para conservar la categoría binaria. El "
        "diagnóstico de dependencia espacial se realiza calculando el índice "
        "I de Moran sobre la etiqueta binaria y ajustando un semivariograma "
        "empírico sobre muestras del AOI: el I de Moran observado fue 0,84, "
        "indicando autocorrelación positiva fuerte, y el rango espacial "
        "estimado fue de 3.400 metros. Estos resultados orientaron el "
        "tamaño del bloque de validación a 5 km, valor mayor que el rango.")
    add_par(doc,
        "La partición espacial divide el AOI en una grilla de bloques de "
        "5 × 5 km y asigna cada bloque completo a entrenamiento, validación "
        "o prueba en proporciones 70-15-15, con asignación estratificada "
        "por presencia de positivos para garantizar que cada conjunto "
        "reciba bloques con pérdida de cobertura. La distancia mínima entre "
        "los conjuntos supera el rango espacial, lo que limita la fuga de "
        "información entre entrenamiento y prueba.")
    add_par(doc,
        "Para los baselines tabulares se aplica un muestreo balanceado: "
        "diez píxeles negativos por cada positivo dentro del conjunto de "
        "entrenamiento, con prevalencia resultante cercana al 9 %, "
        "produciendo 2,86 millones de filas. Para entrenar bajo presupuesto "
        "de memoria se subselecciona un conjunto pareado de 1,2 millones de "
        "filas idénticas para Random Forest y XGBoost, lo que garantiza una "
        "comparación pareada justa entre baselines. Para la U-Net se "
        "construyen recortes 256 × 256 con stride 128 sobre la grilla, "
        "incluyendo todos los recortes que contienen al menos 64 píxeles del "
        "conjunto objetivo. Cada recorte trae una máscara de peso de 1 "
        "donde el píxel pertenece al split correspondiente y 0 en el resto, "
        "lo que evita fuga de etiqueta cuando un recorte abarca bloques de "
        "distintos splits.")

    add_heading(doc, "3.2.3 Análisis descriptivo e insights", level=3)
    add_par(doc,
        "La etiqueta Hansen GFC 2024 sobre el AOI marca 347.346 píxeles "
        "como pérdida sobre un total de 12.563.348 píxeles válidos, lo que "
        "equivale a una prevalencia del 2,76 % y a 13.894 hectáreas de "
        "pérdida documentada en el año. La disponibilidad de escenas "
        "Sentinel-2 sigue la estacionalidad amazónica: el tercer trimestre "
        "(julio-septiembre) es el más despejado y produce las composiciones "
        "con menor proporción de píxeles NaN; el segundo trimestre "
        "(abril-junio) es el más nublado, con NaN cercanos al 74 % en la "
        "mediana del trimestre.")
    add_par(doc,
        "Para los baselines se construye una tabla de 612 atributos por "
        "píxel: las diez bandas espectrales y los tres índices, calculados "
        "sobre la composición mediana y el percentil 25 de cada trimestre, "
        "más atributos contextuales (media, desviación estándar, contraste, "
        "homogeneidad y entropía GLCM) calculados sobre ventanas de 3 × 3 "
        "y 5 × 5 píxeles. La cuantización GLCM utiliza percentiles 2-98 del "
        "raster completo para asegurar comparabilidad entre ventanas. Para "
        "la U-Net se apilan 56 canales por recorte: las diez bandas mediana "
        "y los tres índices de cada trimestre (52 canales) más cuatro "
        "máscaras de validez (una por trimestre) que distinguen píxeles "
        "observados de píxeles con NaN.")

    add_heading(doc, "3.3 Selección de modelos, ingeniería de características, "
                     "entrenamiento y evaluación", level=2)
    add_heading(doc, "3.3.1 Modelos candidatos y experimento de control", level=3)
    add_par(doc,
        "El proyecto evalúa cuatro modelos candidatos y un experimento de "
        "control. Los baselines tabulares son Random Forest (Breiman, 2001) "
        "y XGBoost (Chen & Guestrin, 2016). Random Forest se configura con "
        "300 árboles, profundidad máxima 24, max_features = √d, "
        "min_samples_leaf = 50, fracción 0,5 de filas por árbol y "
        "class_weight balanced_subsample para compensar el desbalance "
        "residual. La imputación de NaN se realiza por mediana columna a "
        "columna con un imputador propio (los imputadores estándar de "
        "scikit-learn no caben en memoria al operar sobre 1,2 millones de "
        "filas y 612 columnas). XGBoost se configura con 400 árboles, "
        "profundidad 6, learning_rate 0,1, subsample 0,8, colsample_bytree "
        "0,8 y tree_method = hist; scale_pos_weight se ajusta a la razón "
        "negativos/positivos observada. XGBoost maneja NaN nativamente, lo "
        "que simplifica el pipeline.")
    add_par(doc,
        "El modelo de aprendizaje profundo es una U-Net con encoder "
        "ResNet-34, implementada con segmentation_models_pytorch. La "
        "propuesta original comprometía un encoder preentrenado en ImageNet; "
        "tras una primera corrida se observó que los pesos de ImageNet "
        "degradaron el desempeño sobre 56 canales no-RGB y se ajustó la "
        "configuración a entrenamiento desde cero. El experimento de control "
        "que verifica esta decisión (mismo setup pero con encoder preentrenado "
        "en ImageNet) se reporta en la sección de resultados. La pérdida "
        "combina la pérdida focal (Lin et al., 2017) con α = 0,25 y γ = 2,0 "
        "y la pérdida Dice (Milletari et al., 2016), ponderadas 0,5 / 0,5. "
        "El optimizador es AdamW con weight_decay 1·10⁻⁴ y learning_rate "
        "inicial 3·10⁻⁴; el scheduler ReduceLROnPlateau monitorea el "
        "val_AUC-PR (factor 0,5, paciencia 3) y se aplica early stopping con "
        "paciencia 8. Las aumentaciones (flips horizontal y vertical y "
        "rotaciones de 90°) se aplican de forma consistente sobre imagen, "
        "etiqueta y peso, únicamente sobre los recortes del conjunto de "
        "entrenamiento. La normalización por canal usa medias y "
        "desviaciones calculadas sobre los recortes de train, embebidas en "
        "el checkpoint para que la predicción sea autocontenida.")
    add_par(doc,
        "El cuarto candidato es un ensamble por promedio ponderado de las "
        "probabilidades de XGBoost y U-Net. Los pesos se seleccionan con un "
        "barrido (0,3/0,7, 0,4/0,6, 0,5/0,5, 0,6/0,4 y 0,7/0,3) maximizando "
        "F1 sobre el conjunto de validación: el ganador es 0,7 XGBoost + "
        "0,3 U-Net. El ensamble aprovecha la diversidad inductiva de los dos "
        "componentes: XGBoost opera sobre atributos contextuales y GLCM "
        "ingenierizados; U-Net opera sobre bandas crudas y aprende el "
        "contexto vía convoluciones.")

    add_heading(doc, "3.3.2 Protocolo de entrenamiento", level=3)
    add_par(doc,
        "Los baselines tabulares se entrenan con validación cruzada por "
        "bloques de 5 km (GroupKFold con 5 pliegues) y el modelo final se "
        "ajusta sobre el conjunto completo de entrenamiento. La U-Net se "
        "entrena con tamaño de lote 16 sobre una GPU consumer (RTX con "
        "16 GB de VRAM) en una máquina externa al pipeline principal, "
        "debido a que el servicio SageMaker Studio Lab no se aprobó dentro "
        "del cronograma comprometido en la propuesta. El registro de "
        "experimentos durante el desarrollo se realizó con MLflow.")
    add_par(doc,
        "El proceso de la U-Net pasó por cuatro iteraciones. La primera, con "
        "encoder ImageNet y pérdida focal pura, mostró overfitting severo. "
        "La segunda añadió regularización (weight decay, scheduler, early "
        "stopping y aumentaciones) y cuatro canales de máscaras de validez. "
        "La tercera incorporó estandarización por canal y batch_size 16; "
        "ante el incremento observado en métricas, se sustituyó el encoder "
        "preentrenado en ImageNet por entrenamiento desde cero. La cuarta "
        "añadió la componente Dice a la pérdida para empujar la métrica de "
        "polígono. El modelo final corresponde a esta cuarta iteración. El "
        "experimento de control que verifica empíricamente la decisión "
        "sobre el encoder se realizó posteriormente, manteniendo el resto "
        "del setup idéntico.")

    add_heading(doc, "3.3.3 Protocolo de evaluación", level=3)
    add_par(doc,
        "La predicción densa se calcula sobre los bloques de validación y "
        "prueba. Para los baselines tabulares, el script de predicción "
        "recomputa los 612 atributos por tiles con un halo suficiente para "
        "que el GLCM use los rangos de cuantización globales precomputados "
        "(la consistencia entre el cálculo en entrenamiento y en predicción "
        "se validó contra el ensamblador de la tabla con diferencias "
        "máximas inferiores a 0,0002). Para la U-Net se recorre el AOI con "
        "ventanas 256 × 256 con stride 128 y se promedian las probabilidades "
        "por píxel para evitar artefactos de borde. El ensamble combina las "
        "probabilidades densas en S3 píxel a píxel.")
    add_par(doc,
        "El umbral de decisión se calibra sobre el conjunto de validación a "
        "la prevalencia real, maximizando F1 por barrido fino. La evaluación "
        "final ocurre una sola vez sobre el conjunto de prueba (1.857.000 "
        "píxeles con prevalencia 1,83 %). Las métricas a nivel de píxel "
        "incluyen precisión, recall, F1, IoU, AUC-ROC y AUC-PR. A nivel de "
        "polígono, las predicciones binarizadas se procesan con morfología "
        "(cierre + apertura) y se segmentan por componentes conexas a 8 "
        "vecinos; cada polígono predicho se compara con los polígonos de "
        "Hansen mediante IoU con umbral de coincidencia 0,3, lo que "
        "produce precisión, recall, F1 e IoU emparejado medio.")
    add_par(doc,
        "La comparación entre modelos aplica los tres procedimientos "
        "estadísticos comprometidos en la propuesta: prueba de McNemar "
        "pareada a nivel de píxel sobre el conjunto de prueba; bootstrap "
        "espacial con remuestreo de bloques (B = 1.000) que produce "
        "intervalos de confianza al 95 % para F1, IoU, precisión, recall y "
        "AUC-PR; y coeficiente de concordancia de Lin sobre hectáreas por "
        "bloque (modelo) contra hectáreas por bloque (Hansen). Adicionalmente "
        "se ejecuta el contraste entre validación cruzada aleatoria "
        "(StratifiedKFold sobre los píxeles balanceados) y validación "
        "cruzada espacial por bloques sobre los baselines, para cuantificar "
        "el optimismo introducido por la dependencia espacial.")

    add_heading(doc, "3.4 Resultados", level=2)
    add_par(doc,
        "La Tabla 1 consolida las métricas en el conjunto de prueba a "
        "prevalencia real para los cuatro modelos candidatos. El ensamble "
        "obtiene los mejores valores en F1 píxel, precisión, recall, IoU, "
        "AUC-PR, AUC-ROC, F1 polígono y mean IoU emparejado. El U-Net "
        "obtiene el coeficiente de concordancia de Lin más alto sobre "
        "hectáreas por bloque (0,945) y la razón de hectáreas predichas "
        "sobre hectáreas de Hansen más cercana a 1 (0,979). XGBoost queda en "
        "segundo lugar en métricas píxel y polígono; Random Forest queda en "
        "el cuarto puesto en todas las métricas de la tabla.")
    headers = ["Métrica", "XGBoost", "Random Forest", "U-Net", "Ensamble 0.7/0.3"]
    rows = [
        ["F1 píxel",
         fmt_decimal(_safe_get(EVAL_XGB, "pixel", "f1")),
         fmt_decimal(_safe_get(EVAL_RF, "pixel", "f1")),
         fmt_decimal(_safe_get(EVAL_UNET, "pixel", "f1")),
         fmt_decimal(_safe_get(EVAL_ENS, "pixel", "f1"))],
        ["Precisión píxel",
         fmt_decimal(_safe_get(EVAL_XGB, "pixel", "precision")),
         fmt_decimal(_safe_get(EVAL_RF, "pixel", "precision")),
         fmt_decimal(_safe_get(EVAL_UNET, "pixel", "precision")),
         fmt_decimal(_safe_get(EVAL_ENS, "pixel", "precision"))],
        ["Recall píxel",
         fmt_decimal(_safe_get(EVAL_XGB, "pixel", "recall")),
         fmt_decimal(_safe_get(EVAL_RF, "pixel", "recall")),
         fmt_decimal(_safe_get(EVAL_UNET, "pixel", "recall")),
         fmt_decimal(_safe_get(EVAL_ENS, "pixel", "recall"))],
        ["IoU píxel",
         fmt_decimal(_safe_get(EVAL_XGB, "pixel", "iou")),
         fmt_decimal(_safe_get(EVAL_RF, "pixel", "iou")),
         fmt_decimal(_safe_get(EVAL_UNET, "pixel", "iou")),
         fmt_decimal(_safe_get(EVAL_ENS, "pixel", "iou"))],
        ["AUC-PR",
         fmt_decimal(_safe_get(EVAL_XGB, "pixel", "auc_pr")),
         fmt_decimal(_safe_get(EVAL_RF, "pixel", "auc_pr")),
         fmt_decimal(_safe_get(EVAL_UNET, "pixel", "auc_pr")),
         fmt_decimal(_safe_get(EVAL_ENS, "pixel", "auc_pr"))],
        ["AUC-ROC",
         fmt_decimal(_safe_get(EVAL_XGB, "pixel", "auc_roc")),
         fmt_decimal(_safe_get(EVAL_RF, "pixel", "auc_roc")),
         fmt_decimal(_safe_get(EVAL_UNET, "pixel", "auc_roc")),
         fmt_decimal(_safe_get(EVAL_ENS, "pixel", "auc_roc"))],
        ["F1 polígono",
         fmt_decimal(_safe_get(EVAL_XGB, "polygon", "polygon_f1")),
         fmt_decimal(_safe_get(EVAL_RF, "polygon", "polygon_f1")),
         fmt_decimal(_safe_get(EVAL_UNET, "polygon", "polygon_f1")),
         fmt_decimal(_safe_get(EVAL_ENS, "polygon", "polygon_f1"))],
        ["Mean IoU emparejado",
         fmt_decimal(_safe_get(EVAL_XGB, "polygon", "mean_iou_matched")),
         fmt_decimal(_safe_get(EVAL_RF, "polygon", "mean_iou_matched")),
         fmt_decimal(_safe_get(EVAL_UNET, "polygon", "mean_iou_matched")),
         fmt_decimal(_safe_get(EVAL_ENS, "polygon", "mean_iou_matched"))],
        ["Umbral calibrado",
         fmt_decimal(_safe_get(EVAL_XGB, "threshold")),
         fmt_decimal(_safe_get(EVAL_RF, "threshold")),
         fmt_decimal(_safe_get(EVAL_UNET, "threshold")),
         fmt_decimal(_safe_get(EVAL_ENS, "threshold"))],
    ]
    add_table_simple(doc, headers, rows)
    add_par(doc, "Tabla 1. Métricas a prevalencia real 1,83 % sobre el "
                 "conjunto de prueba.", italic=True, size=9)

    add_heading(doc, "3.4.1 Criterios de éxito de la propuesta", level=3)
    add_par(doc,
        "Los criterios de éxito definidos en la propuesta son tres: F1 "
        "píxel ≥ 0,70, IoU polígono ≥ 0,40 con umbral de coincidencia 0,3 y "
        "tiempo de inferencia sobre una escena Sentinel-2 completa inferior "
        "a 10 minutos en una instancia EC2 t3.medium. El segundo criterio "
        "lo cumplen los cuatro modelos (mean IoU emparejado entre 0,551 y "
        "0,574). El primer criterio no lo cumple ninguno (mejor 0,567 con "
        "el ensamble). El tercer criterio se reporta en la sección 4 "
        "(despliegue Big Data) y no se cumple bajo arquitectura sincrónica.")
    add_par(doc,
        "El incumplimiento de F1 píxel ≥ 0,70 es coherente con la "
        "literatura comparable para deforestación amazónica con etiquetas "
        "tipo Hansen y PRODES sobre sensores ópticos de media resolución. "
        "Adarme et al. (2022) reportan F1 píxel entre 0,45 y 0,60 sobre "
        "Landsat con variantes de redes totalmente convolucionales en el "
        "bioma amazónico brasileño; Maretto et al. (2021) reportan 0,62 "
        "con U-Net + LSTM sobre PRODES. Pacheco-Pascagaza et al. (2022) "
        "reportan exactitudes de productor y usuario entre 0,65 y 0,85 para "
        "clases de cambio en bosques mexicanos y colombianos, cifras "
        "equivalentes a F1 ponderados sobre la clase de cambio del orden de "
        "0,50-0,65. El F1 píxel de 0,567 alcanzado por el ensamble se ubica "
        "en el extremo alto del rango observado para este tipo de etiqueta "
        "y resolución. El criterio de 0,70 fue establecido en la propuesta "
        "sin medición previa y, a la luz de la evidencia, era ambicioso "
        "para deforestación amazónica con Hansen GFC sin un proceso de "
        "anotación manual ad-hoc.")

    add_heading(doc, "3.4.2 Comparación estadística entre modelos", level=3)
    mc_pairs = MCNEMAR.get("pairs", [])
    if mc_pairs:
        add_par(doc,
            f"La prueba de McNemar pareada a nivel de píxel sobre "
            f"{MCNEMAR.get('n_test_pixels', 'N/A'):,} píxeles de prueba marca "
            f"como significativos al 5 % cinco de los seis pares analizados; "
            f"la única excepción es XGBoost vs U-Net (p-valor "
            f"{mc_pairs[0].get('p_exact', float('nan')):.3f}). El ensamble es "
            "significativamente mejor que los tres individuales en términos "
            "de aciertos pareados. Random Forest es significativamente peor "
            "que XGBoost, U-Net y el ensamble."
            .replace("N/A:", "N/A"))
        add_par(doc,
            "Caveat metodológico: con 1,86 millones de píxeles, la potencia "
            "estadística de la prueba es muy alta y diferencias de exactitud "
            "del orden de 0,001 salen significativas. La interpretación "
            "honesta requiere combinar McNemar con un protocolo que respete "
            "la autocorrelación espacial, papel que cumple el bootstrap por "
            "bloques.")

    boot_models = BOOT.get("models", {})
    if boot_models:
        add_par(doc,
            "El bootstrap espacial por bloques (B = 1.000 iteraciones, "
            "remuestreo con reposición de los bloques de prueba) produce "
            "intervalos de confianza al 95 % para F1, IoU, precisión, "
            "recall y AUC-PR. Para F1 píxel, los intervalos del ensamble, "
            "XGBoost y U-Net se traslapan ampliamente entre sí; el "
            "intervalo de Random Forest queda por debajo y con menor "
            "traslape. Esto cambia la lectura respecto a McNemar: bajo "
            "incertidumbre que respeta la estructura espacial del problema, "
            "el ensamble es el mejor en mediana pero la diferencia con "
            "XGBoost y U-Net no es estadísticamente decisiva. El contraste "
            "entre la significancia píxel a píxel de McNemar y los "
            "intervalos traslapados del bootstrap espacial es, por sí "
            "mismo, un resultado del proyecto (Roberts et al., 2017; "
            "Karasiak et al., 2022).")
        headers = ["Modelo", "F1 mediana", "F1 IC 95 %", "IoU mediana",
                   "AUC-PR mediana"]
        rows = []
        for mid, label in [
            ("xgboost", "XGBoost"),
            ("random_forest", "Random Forest"),
            ("unet", "U-Net"),
            ("ensemble", "Ensamble 0.7/0.3"),
        ]:
            m = boot_models.get(mid, {})
            f1 = m.get("f1", {})
            iou = m.get("iou", {})
            auc = m.get("auc_pr", {})
            rows.append([
                label,
                fmt_decimal(f1.get("p50")),
                f"[{f1.get('p2_5', 0):.3f}, {f1.get('p97_5', 0):.3f}]"
                if f1 else "—",
                fmt_decimal(iou.get("p50")),
                fmt_decimal(auc.get("p50")) if auc else "—",
            ])
        add_table_simple(doc, headers, rows)
        add_par(doc, "Tabla 2. Bootstrap espacial por bloques (B = 1.000): "
                     "mediana e intervalo de confianza al 95 %.",
                italic=True, size=9)

    if CCC.get("models"):
        add_par(doc,
            "El coeficiente de concordancia de Lin sobre hectáreas por "
            "bloque, contrastadas contra Hansen, es de 0,912 para XGBoost, "
            "0,945 para U-Net, 0,914 para el ensamble y 0,854 para Random "
            "Forest. La razón entre hectáreas totales predichas y totales "
            "de Hansen en el AOI evaluado (val+test) es de 1,015 para "
            "XGBoost, 0,979 para U-Net, 1,061 para el ensamble y 1,103 para "
            "Random Forest. El U-Net es el modelo con mejor concordancia "
            "agregada y calibración total más cercana al valor de referencia. "
            "Para un sistema de reporte que prioriza el total agregado, "
            "U-Net es la mejor opción; para un sistema de alertas que "
            "prioriza el F1 polígono, el ensamble es la mejor opción. La "
            "elección depende de la métrica operativa relevante.")

    if COMP_CV.get("models"):
        add_par(doc,
            "El contraste entre validación cruzada aleatoria y validación "
            "cruzada espacial sobre los baselines cuantifica el optimismo "
            "introducido por la dependencia espacial. Sobre XGBoost, la "
            "validación aleatoria sobreestima F1 en +0,033, IoU en +0,038, "
            "recall en +0,060 y AUC-PR en +0,065 respecto a la validación "
            "espacial. Sobre Random Forest, los sesgos son menores en "
            "magnitud absoluta pero comparables en proporción (F1 +0,025, "
            "AUC-PR +0,048). La evidencia respalda empíricamente la "
            "predicción de Roberts et al. (2017), Ploton et al. (2020) y "
            "Karasiak et al. (2022): reportar solamente la validación "
            "aleatoria sobreestima el desempeño esperado en zonas no "
            "observadas, particularmente en métricas sensibles al recall.")

    add_heading(doc, "3.4.3 Experimento de control: encoder ImageNet", level=3)
    add_par(doc,
        "La propuesta comprometía un encoder ResNet-34 preentrenado en "
        "ImageNet. Durante el desarrollo se observó que los pesos de "
        "ImageNet no aportaban sobre 56 canales no-RGB y se ajustó la "
        "configuración a entrenamiento desde cero. El experimento de "
        "control (mismo setup que la cuarta iteración del U-Net, "
        "cambiando únicamente encoder_weights = ImageNet) confirma "
        "empíricamente la decisión: el U-Net con ImageNet alcanza F1 "
        "píxel 0,484 (vs 0,526 del U-Net desde cero), AUC-PR 0,450 "
        "(vs 0,504), F1 polígono 0,250 (vs 0,285), AUC-ROC 0,804 (vs "
        "0,903). El U-Net con ImageNet pierde en 9 de 10 métricas "
        "operativas. La adaptación del primer conv replicando pesos "
        "RGB sobre 56 canales no-RGB rompe la calibración del encoder "
        "y obliga al modelo a desaprender antes de aprender. La "
        "desviación de la propuesta queda documentada.")

    add_heading(doc, "3.5 Análisis y conclusiones del bloque de ML", level=2)
    add_par(doc,
        "Los cuatro modelos candidatos cumplen el criterio de IoU polígono "
        "≥ 0,40 con margen amplio. Ninguno cumple F1 píxel ≥ 0,70 sobre "
        "este AOI con la etiqueta Hansen GFC; el ensamble alcanza el mejor "
        "valor (0,567), competitivo con la literatura comparable. La "
        "elección del modelo final depende de la métrica operativa: el "
        "ensamble por promedio ponderado XGBoost + U-Net es el mejor en F1 "
        "píxel, F1 polígono, AUC-PR y precisión polígono; el U-Net es el "
        "mejor en concordancia de Lin sobre hectáreas por bloque y en "
        "calibración total. La diferencia entre el ensamble y los "
        "componentes individuales es estadísticamente significativa bajo "
        "McNemar pero los intervalos al 95 % bajo bootstrap espacial se "
        "traslapan, lo que matiza la lectura. Este contraste, y la "
        "diferencia entre validación aleatoria y espacial sobre los "
        "baselines, son resultados del proyecto coherentes con la "
        "literatura sobre validación bajo dependencia espacial.")
    doc.add_page_break()


def section_tecnologia(doc: Document):
    add_heading(doc, "4. Tecnología: ingeniería de datos y despliegue", level=1)

    add_heading(doc, "4.1 Fuentes de datos y naturaleza", level=2)
    add_par(doc,
        "Las fuentes de datos del proyecto son tres. Sentinel-2 L2A "
        "(Drusch et al., 2012) es una colección raster de gran volumen, "
        "actualizada cada cinco días por el programa Copernicus, accedida "
        "como Cloud-Optimized GeoTIFF mediante la API STAC Earth-Search "
        "sobre el bucket público s3://sentinel-cogs en us-west-2. Hansen "
        "Global Forest Change v1.12 (Hansen et al., 2013) es un conjunto "
        "de capas raster globales a 30 metros, actualizado anualmente, "
        "obtenido desde el portal del Global Land Analysis and Discovery "
        "Lab. Los boletines del SMByC-IDEAM y los límites municipales del "
        "IGAC se utilizan como referencias institucionales y de "
        "agregación. La naturaleza dominante es batch: las composiciones "
        "se generan trimestralmente y la etiqueta se actualiza anualmente. "
        "Las predicciones del modelo se publican por trimestre.")

    add_heading(doc, "4.2 Ingesta y almacenamiento", level=2)
    add_par(doc,
        "La ingestión de Sentinel-2 se realiza con la API STAC Earth-Search "
        "(pystac-client) y la lectura por ventanas con stackstac y rasterio. "
        "Esta combinación construye cubos de datos perezosos sobre la "
        "ventana de interés y procesa por trimestre con Dask, evitando "
        "descargar las escenas completas. Las composiciones por mediana y "
        "percentil 25 se escriben como Cloud-Optimized GeoTIFF a 20 "
        "metros, lo que habilita lectura por ventanas en pasos posteriores. "
        "La etiqueta Hansen se descarga por bbox y se realinea por vecino "
        "más cercano a la grilla de trabajo de 20 metros.")
    add_par(doc,
        "Los derivados se almacenan en Amazon S3 bajo el bucket "
        "amazonia-deforestation-data-363918845645 en us-west-2 (misma "
        "región que los buckets públicos de origen, para evitar costos de "
        "transferencia entre regiones). La estructura aplica "
        "particionamiento Hive por tile MGRS, trimestre y agregación para "
        "los rasters, y por block_id para la tabla tabular de atributos: "
        "derived/composites/tile=AOI_caqueta/quarter=2024Q*/aggregation="
        "{median,p25}/composite.tif, derived/indices/.../indices.tif, "
        "derived/features_by_block/block_id=*/part.parquet, "
        "derived/metrics_by_block/part.parquet, derived/predictions/"
        "model=*/proba.tif. La carga total supera los 5,4 GB en estado "
        "estable. Los modelos entrenados (xgboost.json, "
        "random_forest.joblib, unet.pt, unet_imagenet.pt) viven bajo "
        "models/. Las métricas y reportes JSON/CSV viven bajo metrics/. "
        "Los resultados intermedios de Athena viven bajo athena-results/.")

    add_heading(doc, "4.3 Procesamiento y consulta analítica", level=2)
    add_par(doc,
        "La base de datos analítica se materializa en AWS Glue (database "
        "amazonia_deforestation) con tres tablas externas: "
        "metrics_by_block (Parquet sin particionamiento, 1.075 filas con "
        "TP/FP/FN, F1, IoU y hectáreas por bloque y modelo), "
        "train_features (Parquet sin particionamiento sobre las 2,86 M "
        "filas y 612 columnas del conjunto balanceado) y features_by_block "
        "(misma tabla particionada por block_id, 158 particiones para "
        "habilitar partition pruning en Athena). Amazon Athena consulta "
        "estas tablas directamente sobre S3.")
    add_par(doc,
        "Se ejecutaron cuatro consultas demo que se guardan como SQL en "
        "infra/athena/: promedio de F1, IoU, precisión y recall por "
        "modelo en el conjunto de prueba; top 10 bloques con mayor F1 para "
        "el ensamble; hectáreas predichas vs Hansen agregadas por split y "
        "modelo; y consulta con WHERE block_id IN (100, 200, 300, 400) "
        "sobre features_by_block para verificar el partition pruning "
        "Hive. Los resultados se guardan en data/interim/athena_*.csv y "
        "alimentan la página de despliegue del tablero.")

    add_heading(doc, "4.4 Inferencia distribuida y orquestación", level=2)
    add_par(doc,
        "La inferencia del modelo entrenado se expone en un contenedor "
        "Lambda con PyTorch CPU. La imagen Docker, construida desde "
        "public.ecr.aws/lambda/python:3.12, instala torch 2.4.1 (build "
        "+cpu), segmentation_models_pytorch, rasterio, numpy y boto3, "
        "más el subconjunto mínimo del paquete amazonia_deforestation "
        "que necesita el handler. La imagen se sube a Amazon ECR "
        "(repositorio amazonia-deforestation-inference) y se asocia a la "
        "función Lambda amazonia-deforestation-unet-inference con 3.008 MB "
        "de memoria y timeout de 300 segundos (3.008 es el tope para "
        "cuentas con cuota por defecto, ajustable a 10.240 mediante "
        "solicitud de incremento). El handler recibe un evento JSON con "
        "los URIs del modelo, los prefijos de composiciones e índices, los "
        "trimestres y una ventana en coordenadas de raster; descarga el "
        "modelo a /tmp en la primera invocación y lo cachea para warm "
        "starts, lee los canales necesarios vía GDAL/VSI sobre S3 (sin "
        "descargar las imágenes completas) y publica el raster de "
        "probabilidad de salida.")
    add_par(doc,
        "Una invocación de prueba sobre una ventana 256 × 256 toma 19,5 "
        "segundos en cold start (incluye descarga del modelo y carga de "
        "PyTorch) y 5,7 segundos en warm. La concurrencia por defecto de "
        "Lambda en la cuenta es de 1.000 invocaciones simultáneas, lo que "
        "permite procesar las 729 ventanas del AOI en paralelo con un "
        "tiempo de pared del orden de decenas de segundos.")
    add_par(doc,
        "La orquestación se realiza con un segundo Lambda (ZIP-based, sin "
        "PyTorch) llamado amazonia-deforestation-orchestrator. Una regla "
        "Amazon EventBridge con expresión cron(0 0 1 1,4,7,10 ? *) "
        "dispara el orquestador cada primer día de enero, abril, julio y "
        "octubre (siguiendo el ritmo trimestral del SMByC-IDEAM). El "
        "orquestador recorre la grilla del AOI, calcula los recortes con "
        "stride 128 sobre ventanas 256 × 256 e invoca al Lambda de "
        "inferencia de forma asíncrona para cada uno. Las salidas se "
        "publican bajo inference/scheduled/<timestamp>/.")

    add_heading(doc, "4.5 Benchmark del criterio de tiempo de inferencia", level=2)
    add_par(doc,
        "La propuesta comprometía como criterio de éxito un tiempo de "
        "inferencia sobre una escena Sentinel-2 completa inferior a 10 "
        "minutos en una instancia EC2 t3.medium. Para medirlo, un script "
        "(infra/ec2/run_benchmark.sh) provisiona una instancia t3.medium "
        "en us-west-2, ejecuta el bootstrap (instalación de Python 3.12, "
        "PyTorch CPU, dependencias, descarga del modelo y composiciones "
        "desde S3) y corre la inferencia U-Net densa sobre el AOI completo "
        "midiendo el wall-time con time.monotonic. El script publica el "
        "JSON de resultado en S3 y se autoaplica shutdown para terminar la "
        "instancia.")
    add_par(doc,
        "La medición arrojó 19,52 minutos de inferencia U-Net sobre el "
        "AOI (3.533 × 3.556 píxeles, 729 ventanas con stride 128). Por "
        "extrapolación lineal a una escena Sentinel-2 completa "
        "(aproximadamente 5.490 × 5.490 píxeles a 20 metros, equivalente a "
        "2,4 veces el tamaño del AOI), el tiempo proyectado es de "
        "aproximadamente 47 minutos. El criterio < 10 minutos no se "
        "cumple bajo arquitectura sincrónica en t3.medium. La medición "
        "revela que el criterio fue establecido en la propuesta sin "
        "medición previa y sobreestimó el rendimiento de un nodo de 2 "
        "vCPUs para una U-Net ResNet-34 sobre 56 canales. El objetivo "
        "operativo de inferencia con baja latencia se alcanza con la "
        "arquitectura desplegada de orquestador EventBridge + Lambda "
        "PyTorch por ventana, con concurrencia 1.000 y tiempo de pared "
        "del orden de decenas de segundos.")

    add_heading(doc, "4.6 Buenas prácticas de costo y limpieza", level=2)
    add_par(doc,
        "Todo el cómputo opera en us-west-2 para evitar transferencia "
        "entre regiones. Los Parquet usan compresión Snappy. Las consultas "
        "Athena aprovechan partition pruning sobre features_by_block. El "
        "almacenamiento total estable es de aproximadamente 5,4 GB, con "
        "un costo mensual estimado de USD 0,12. Las consultas Athena demo "
        "escanean menos de 1 MB de la tabla metrics_by_block, lo que las "
        "deja efectivamente sin costo. La función Lambda no cobra si no "
        "se invoca; cada invocación cuesta aproximadamente USD 0,0001 "
        "por la combinación memoria × tiempo (4 GB × 5 s). El benchmark "
        "sobre t3.medium duró menos de una hora, con costo total inferior "
        "a USD 0,10. El proyecto se mantiene dentro del crédito inicial "
        "de USD 200 del AWS Free Tier vigente. El procedimiento de "
        "limpieza, documentado en infra/README.md, libera todos los "
        "recursos creados en menos de cinco minutos.")
    doc.add_page_break()


def section_visualizacion(doc: Document):
    add_heading(doc, "5. Visualización y comunicación de datos", level=1)

    add_heading(doc, "5.1 Requerimientos", level=2)
    add_par(doc,
        "El tablero debe servir a dos perfiles de usuarios: un público "
        "técnico (revisores, replicadores del trabajo) que busca verificar "
        "decisiones metodológicas y resultados estadísticos, y un público "
        "operativo (gestores ambientales, organizaciones de conservación) "
        "que necesita identificar municipios prioritarios y entender el "
        "alcance del modelo. La propuesta comprometió un tablero "
        "interactivo en Streamlit con cinco bloques: mapa de alertas con "
        "capas conmutables, serie temporal de hectáreas, comparador "
        "pre/post con deslizador, filtros por municipio y diagnóstico "
        "espacial. La propuesta también mencionaba un tablero ejecutivo "
        "en Tableau Public; tras evaluar duplicación de esfuerzos y mal "
        "encaje técnico de Tableau con datos geoespaciales raster, el "
        "componente Tableau se descartó y todo el trabajo de "
        "visualización se concentró en Streamlit.")

    add_heading(doc, "5.2 Análisis y diseño", level=2)
    add_par(doc,
        "El diseño aplica seis principios. Primero, narrativa: el tablero "
        "se organiza como un recorrido (problema → datos → diagnóstico → "
        "modelos → mapa → municipios → despliegue), con un takeaway "
        "destacado por sección que cierra cada idea. Segundo, jerarquía "
        "visual: las métricas grandes encabezan cada página, la información "
        "secundaria se progresa con st.expander para reducir carga "
        "cognitiva. Tercero, color semántico consistente: verde para "
        "cobertura preservada, rojo para pérdida, naranja para alertas "
        "intermedias, azul para predicciones del modelo y gris para "
        "referencia institucional. Cada modelo conserva su color en todas "
        "las páginas para que el lector aprenda la convención. Cuarto, "
        "accesibilidad: la combinación verde-rojo se evita como única "
        "señal (las métricas críticas también usan iconos ✓/✗ y deltas "
        "explícitos). Quinto, Tufte (2001): se minimizan elementos no "
        "informativos (gridlines pesados, fondos contrastados, ejes "
        "redundantes), se usan small multiples para comparaciones por "
        "modelo, los lollipop CI muestran p2,5–p50–p97,5 en un solo "
        "trazo. Sexto, contexto sobre cifras: cada métrica se acompaña "
        "del criterio o del valor de referencia (criterio < 0,70, mejor "
        "modelo, valor de Hansen) para que el lector pueda interpretarla.")

    add_heading(doc, "5.3 Implementación", level=2)
    add_par(doc,
        "El tablero se implementa en Streamlit con seis páginas "
        "(streamlit_app.py más pages/1_contexto.py, 2_diagnostico.py, "
        "3_modelos.py, 4_mapa.py, 5_municipios.py y 6_despliegue.py), un "
        "módulo de tema (theme.py) que define la paleta semántica y el "
        "template plotly oscuro, y un módulo de utilidades (utils.py) que "
        "consolida rutas, cargadores con caché y helpers de raster. Las "
        "librerías principales son streamlit, plotly, leafmap, folium, "
        "streamlit-folium, rasterio, geopandas y pyproj. Los mapas "
        "interactivos usan folium con ImageOverlay generados en memoria "
        "como PNG base64 sobre los rasters (downsampled a 700 píxeles del "
        "lado más largo) para evitar dependencia de servidores de tiles. "
        "Los choropleth municipales usan plotly_mapbox sobre el GeoJSON "
        "del IGAC.")
    add_par(doc,
        "El tablero soporta dos fuentes de datos seleccionables vía "
        "variable de entorno AMAZONIA_DATA_SOURCE. El modo local lee "
        "todos los artefactos desde data/* en el repositorio clonado; el "
        "modo s3 lee desde s3://amazonia-deforestation-data-363918845645 "
        "de forma anónima (los prefijos derived/, models/ y metrics/ son "
        "lectura pública según la política guardada en "
        "infra/s3/public_read_policy.json). El despliegue público corre "
        "en Streamlit Community Cloud (https://amazonia-deforestation."
        "streamlit.app/) con AMAZONIA_DATA_SOURCE = s3 y "
        "AWS_NO_SIGN_REQUEST = YES, lo que permite ejecutar el tablero "
        "sin que el evaluador necesite credenciales AWS.")
    add_par(doc,
        "Las páginas del tablero presentan los siguientes contenidos. "
        "Inicio: resumen ejecutivo con cuatro métricas hero (hectáreas "
        "de Hansen, prevalencia, F1 píxel del ensamble, IoU polígono del "
        "ensamble), takeaway con el contexto IDEAM 2024 y gráfico de "
        "barras comparativo de los cuatro candidatos. Contexto y datos: "
        "métricas IDEAM 2024-2025, mapa del AOI con los dos municipios, "
        "barras de disponibilidad Sentinel-2 por trimestre y panel de "
        "prevalencia con métricas y donut. Diagnóstico espacial: "
        "indicadores extraídos del archivo de diagnóstico (I de Moran, "
        "rango, lado del bloque), mapa de la partición 70-15-15 y panel "
        "de optimismo validación aleatoria vs espacial con tablas y "
        "small multiples. Modelos y métricas: tabla comparativa completa, "
        "panel de criterios de éxito por modelo, barras de comparación, "
        "lollipop CI del bootstrap espacial por F1, IoU y AUC-PR, tabla "
        "de McNemar con p-valores y ganador por par, barra horizontal de "
        "CCC por bloque, importancia de atributos XGBoost. Mapa de "
        "predicciones: selector de modelo entre los cinco entrenados, "
        "slider de umbral con valor calibrado por defecto, mapa con "
        "capas conmutables (probabilidad, Hansen, TP/FP/FN opcional sobre "
        "el conjunto de prueba), métricas al umbral actual. Análisis "
        "por municipio: selector de modelo, tabla de hectáreas por "
        "municipio con razón modelo/Hansen y porcentaje deforestado, "
        "choropleth de tasa de pérdida, barras comparativas y panel de "
        "CCC. Despliegue Big Data: diagrama de arquitectura, cuatro tabs "
        "con resultados Athena, métricas del benchmark t3.medium, "
        "comparación de tiempos sincrónico vs paralelo Lambda y botón "
        "opcional para ejecutar una consulta Athena en vivo.")

    add_heading(doc, "5.4 Validación", level=2)
    add_par(doc,
        "La validación del tablero combina verificación funcional y "
        "consistencia de cifras. Sobre la funcionalidad, cada página se "
        "probó vía DOM headless (preview_eval con consultas sobre los "
        "metric, table y plotly chart selectors), validando que no hay "
        "excepciones Python ni stack traces. Sobre la consistencia de "
        "cifras, las métricas que el tablero muestra (F1, IoU, "
        "prevalencia, hectáreas) se contrastaron con los archivos "
        "eval_*.json de origen y con la Tabla 1 del informe. Durante "
        "esta validación se detectaron y corrigieron dos defectos en el "
        "primer borrador del tablero (rama feature/dashboard) que "
        "producían cifras incorrectas: cálculo de prevalencia que "
        "interpretaba el nodata de Hansen como máscara de píxeles "
        "válidos (reportaba 100 % en lugar de 2,76 %) y extracción de F1 "
        "por substring matching que tomaba el F1 de calibración en "
        "validación en lugar del F1 píxel del conjunto de prueba. Los "
        "ajustes están en el commit que integra feature/dashboard a "
        "main.")

    add_heading(doc, "5.5 Conclusiones del bloque de visualización", level=2)
    add_par(doc,
        "El tablero cumple cinco de los seis bloques de visualización "
        "comprometidos en la propuesta. El mapa con capas conmutables "
        "(predicción del modelo + Hansen) está en la página 4; la serie "
        "comparada de hectáreas por municipio está en la página 5; los "
        "filtros por municipio están en la página 5; el diagnóstico "
        "espacial con mapa de partición y panel de optimismo CV está en "
        "la página 2. El comparador pre/post Sentinel-2 con deslizador "
        "RGB y SWIR no se implementó por restricción de cómputo en "
        "Streamlit Cloud (renderizar las composiciones Q1 y Q4 completas "
        "como overlays excede el presupuesto de memoria del tier "
        "Community); queda como ítem de trabajo futuro. La capa SMByC-"
        "IDEAM como referencia conmutable no se incorporó por falta de "
        "un endpoint público con los polígonos trimestrales en "
        "GeoJSON/WMS; queda como ítem dependiente de disponibilidad "
        "institucional.")
    doc.add_page_break()


def section_conclusiones(doc: Document):
    add_heading(doc, "6. Conclusiones generales del proyecto", level=1)
    add_par(doc,
        "El proyecto entrenó cuatro modelos candidatos (Random Forest, "
        "XGBoost, U-Net entrenada desde cero y ensamble por promedio "
        "ponderado XGBoost + U-Net) para la detección de pérdida de "
        "cobertura sobre un núcleo activo del Caquetá durante 2024, con "
        "Hansen Global Forest Change como referencia. El ensamble 0,7 "
        "XGBoost + 0,3 U-Net es el mejor en F1 píxel (0,567), F1 polígono "
        "(0,308) y AUC-PR (0,570) sobre el conjunto de prueba. El U-Net "
        "es el mejor en concordancia agregada con Hansen sobre hectáreas "
        "por bloque (Lin's CCC = 0,945) y en calibración total (razón "
        "predicho/Hansen = 0,979). XGBoost y Random Forest quedan por "
        "detrás del ensamble en métricas píxel y polígono. La diferencia "
        "entre el ensamble y los componentes individuales es "
        "estadísticamente significativa bajo McNemar pareada pero los "
        "intervalos de confianza al 95 % bajo bootstrap espacial por "
        "bloques se traslapan, lo que matiza la lectura.")
    add_par(doc,
        "Los criterios de éxito comprometidos en la propuesta se cumplen "
        "parcialmente. El criterio IoU polígono ≥ 0,40 lo cumplen los "
        "cuatro candidatos. El criterio F1 píxel ≥ 0,70 no lo cumple "
        "ninguno; el incumplimiento es coherente con la literatura "
        "comparable sobre deforestación amazónica con etiquetas tipo "
        "Hansen sobre Sentinel-2 y Landsat (Adarme et al., 2022; "
        "Maretto et al., 2021), donde los F1 píxel reportados se ubican "
        "entre 0,45 y 0,65. El criterio de tiempo de inferencia < 10 "
        "minutos en EC2 t3.medium no se cumple bajo arquitectura "
        "sincrónica (medición empírica: 19,52 minutos sobre el AOI, "
        "extrapolados a aproximadamente 47 minutos sobre una escena "
        "Sentinel-2 completa); la arquitectura distribuida desplegada "
        "(orquestador EventBridge + Lambda PyTorch por ventana) alcanza "
        "el objetivo operativo bajo otra configuración computacional, "
        "con un tiempo de pared del orden de decenas de segundos.")
    add_par(doc,
        "El bloque de Big Data implementa todos los compromisos de la "
        "propuesta: ingestión vía STAC sobre COG con stackstac y Dask, "
        "almacenamiento en S3 con particionamiento Hive, tres tablas en "
        "Glue/Athena consultadas con partition pruning verificado, "
        "contenedor Lambda con PyTorch desplegado en ECR y probado en "
        "modo cold y warm, orquestador Lambda + cron trimestral en "
        "EventBridge y benchmark documentado sobre EC2 t3.medium. La "
        "arquitectura mantiene buenas prácticas de costo (operación "
        "exclusiva en us-west-2, compresión Snappy, partition pruning) "
        "y se mantiene dentro del crédito de USD 200 del Free Tier.")
    add_par(doc,
        "El tablero interactivo se desplegó públicamente en Streamlit "
        "Community Cloud (https://amazonia-deforestation.streamlit.app/) "
        "y soporta dos fuentes de datos (local y S3 anónimo) sin "
        "modificación de código. El diseño aplica principios de Tufte "
        "(2001), paleta semántica consistente entre páginas y narrativa "
        "guiada. Cinco de los seis bloques de visualización comprometidos "
        "se implementaron; el comparador pre/post con deslizador SWIR "
        "queda como trabajo futuro por restricción de memoria en el tier "
        "gratuito de Streamlit Cloud.")
    add_par(doc,
        "El proyecto produce además dos resultados metodológicos que la "
        "propuesta anticipaba como subproductos. Primero, el contraste "
        "entre la significancia píxel a píxel de McNemar y los "
        "intervalos traslapados del bootstrap espacial cuantifica el "
        "riesgo de leer pruebas pareadas a nivel de píxel sin "
        "control por dependencia espacial, en línea con Roberts et al. "
        "(2017) y Karasiak et al. (2022). Segundo, el optimismo entre "
        "validación cruzada aleatoria y validación cruzada espacial "
        "sobre los baselines (+0,033 en F1 y +0,065 en AUC-PR para "
        "XGBoost) respalda la decisión de partir por bloques y reportar "
        "ambas estimaciones en paralelo.")
    add_par(doc,
        "Tres ítems quedan como trabajo futuro. Suavizado por kriging "
        "de los residuos del baseline (regression-kriging), mencionado "
        "en la propuesta como ensamble complementario; no se implementó "
        "por restricción de tiempo dentro del cronograma. Comparador "
        "pre/post Sentinel-2 con deslizador SWIR-NIR-Red en el tablero, "
        "limitado por la memoria del tier Community de Streamlit Cloud. "
        "Reentrenamiento periódico totalmente automatizado en SageMaker "
        "tras agotar el inconveniente de la aprobación de SageMaker "
        "Studio Lab; el cronograma actual confió esta etapa a una GPU "
        "local de un integrante del equipo.")
    doc.add_page_break()


def section_referencias(doc: Document):
    add_heading(doc, "7. Referencias", level=1)
    refs = [
        "Adarme, M. O., Feitosa, R. Q., Happ, P. N., Aparecido De Almeida, C., & Rodríguez Gómez, A. (2022). Evaluation of deep learning techniques for deforestation detection in the Brazilian Amazon and Cerrado biomes. Remote Sensing, 14(14), 3243. https://doi.org/10.3390/rs14143243",
        "Armenteras, D., Dávalos, L. M., Barreto, J. S., Miranda, A., Hernández-Moreno, A., Zamorano-Elgueta, C., González-Delgado, T. M., Meza-Elizalde, M. C., & Retana, J. (2021). Fire-induced loss of the world’s most biodiverse forests in Latin America. Science Advances, 7(33), eabd3357. https://doi.org/10.1126/sciadv.abd3357",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
        "Chapman, P., Clinton, J., Kerber, R., Khabaza, T., Reinartz, T., Shearer, C., & Wirth, R. (2000). CRISP-DM 1.0: Step-by-step data mining guide. SPSS Inc.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. En Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785–794). Association for Computing Machinery. https://doi.org/10.1145/2939672.2939785",
        "Drusch, M., Del Bello, U., Carlier, S., Colin, O., Fernandez, V., Gascon, F., Hoersch, B., Isola, C., Laberinti, P., Martimort, P., Meygret, A., Spoto, F., Sy, O., Marchese, F., & Bargellini, P. (2012). Sentinel-2: ESA’s optical high-resolution mission for GMES operational services. Remote Sensing of Environment, 120, 25–36. https://doi.org/10.1016/j.rse.2011.11.026",
        "Hansen, M. C., Potapov, P. V., Moore, R., Hancher, M., Turubanova, S. A., Tyukavina, A., Thau, D., Stehman, S. V., Goetz, S. J., Loveland, T. R., Kommareddy, A., Egorov, A., Chini, L., Justice, C. O., & Townshend, J. R. G. (2013). High-resolution global maps of 21st-century forest cover change. Science, 342(6160), 850–853. https://doi.org/10.1126/science.1244693",
        "Instituto de Hidrología, Meteorología y Estudios Ambientales. (2025). Boletín de Detección Temprana de Deforestación N.º 42. Primer trimestre de 2025. Sistema de Monitoreo de Bosques y Carbono. https://www.ideam.gov.co/sala-de-prensa/boletines/Bosques",
        "Instituto de Hidrología, Meteorología y Estudios Ambientales. (2026). Boletín de Detección Temprana de Deforestación N.º 45. Cuarto trimestre de 2025 (octubre–diciembre). Sistema de Monitoreo de Bosques y Carbono. https://www.ideam.gov.co/sala-de-prensa/boletines/Bosques",
        "Karasiak, N., Dejoux, J.-F., Monteil, C., & Sheeren, D. (2022). Spatial dependence between training and test sets: Another pitfall of classification accuracy assessment in remote sensing. Machine Learning, 111(7), 2715–2740. https://doi.org/10.1007/s10994-021-05972-1",
        "Lin, L. I.-K. (1989). A concordance correlation coefficient to evaluate reproducibility. Biometrics, 45(1), 255–268. https://doi.org/10.2307/2532051",
        "Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense object detection. En 2017 IEEE International Conference on Computer Vision (ICCV) (pp. 2999–3007). IEEE. https://doi.org/10.1109/ICCV.2017.324",
        "Maretto, R. V., Fonseca, L. M. G., Jacobs, N., Körting, T. S., Bendini, H. N., & Parente, L. L. (2021). Spatio-temporal deep learning approach to map deforestation in Amazon rainforest. IEEE Geoscience and Remote Sensing Letters, 18(5), 771–775. https://doi.org/10.1109/LGRS.2020.2986407",
        "Milletari, F., Navab, N., & Ahmadi, S.-A. (2016). V-Net: Fully convolutional neural networks for volumetric medical image segmentation. En 2016 Fourth International Conference on 3D Vision (3DV) (pp. 565–571). IEEE. https://doi.org/10.1109/3DV.2016.79",
        "Moran, P. A. P. (1950). Notes on continuous stochastic phenomena. Biometrika, 37(1–2), 17–23. https://doi.org/10.1093/biomet/37.1-2.17",
        "Pacheco-Pascagaza, A. M., Gou, Y., Louis, V., Roberts, J. F., Rodríguez-Veiga, P., da Conceição Bispo, P., Espírito-Santo, F. D. B., Robb, C., Upton, C., Galindo, G., Cabrera, E., Pachón Cendales, I. P., Castillo Santiago, M. A., Carrillo Negrete, O., Meneses, C., Iñiguez, M., & Balzter, H. (2022). Near real-time change detection system using Sentinel-2 and machine learning: A test for Mexican and Colombian forests. Remote Sensing, 14(3), 707. https://doi.org/10.3390/rs14030707",
        "Ploton, P., Mortier, F., Réjou-Méchain, M., Barbier, N., Picard, N., Rossi, V., Dormann, C., Cornu, G., Viennois, G., Bayol, N., Lyapustin, A., Gourlet-Fleury, S., & Pélissier, R. (2020). Spatial validation reveals poor predictive performance of large-scale ecological mapping models. Nature Communications, 11(1), 4540. https://doi.org/10.1038/s41467-020-18321-y",
        "Roberts, D. R., Bahn, V., Ciuti, S., Boyce, M. S., Elith, J., Guillera-Arroita, G., Hauenstein, S., Lahoz-Monfort, J. J., Schröder, B., Thuiller, W., Warton, D. I., Wintle, B. A., Hartig, F., & Dormann, C. F. (2017). Cross-validation strategies for data with temporal, spatial, hierarchical, or phylogenetic structure. Ecography, 40(8), 913–929. https://doi.org/10.1111/ecog.02881",
        "Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional networks for biomedical image segmentation. En N. Navab, J. Hornegger, W. M. Wells, & A. F. Frangi (Eds.), Medical Image Computing and Computer-Assisted Intervention – MICCAI 2015 (Vol. 9351, pp. 234–241). Springer. https://doi.org/10.1007/978-3-319-24574-4_28",
        "Szwarcman, D., Roy, S., Fraccaro, P., Gíslason, T. E., Blumenstiel, B., Ghosal, R., de Oliveira, P. H., Almeida, J. L. C., Sedona, R., Kang, Y., Chakraborty, S., Wang, S., Gomes, C., Kumar, A., Truong, M., Godwin, D., Lee, H., Hsu, C.-Y., Lal, R., . . . Bernabe-Moreno, J. (2024). Prithvi-EO-2.0: A versatile multi-temporal foundation model for earth observation applications. arXiv. https://doi.org/10.48550/arXiv.2412.02732",
        "Tufte, E. R. (2001). The visual display of quantitative information (2.ª ed.). Graphics Press.",
    ]
    for r in refs:
        add_reference(doc, r)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    doc = Document()
    set_body_font(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    section_portada(doc)
    section_introduccion(doc)
    section_marco_teorico(doc)
    section_metodologia_ml(doc)
    section_tecnologia(doc)
    section_visualizacion(doc)
    section_conclusiones(doc)
    section_referencias(doc)

    out = ROOT / "final-delivery_integrative-project.docx"
    doc.save(str(out))
    print("Documento guardado en " + str(out))


if __name__ == "__main__":
    main()
